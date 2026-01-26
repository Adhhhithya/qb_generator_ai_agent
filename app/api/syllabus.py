import io
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
from typing import Optional
import logging

# PDF extraction
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

# DOCX extraction
try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/syllabus", tags=["Syllabus Management"])


class SyllabusExtractResponse(BaseModel):
    """Response model for syllabus extraction"""
    raw_text: str
    file_name: str
    file_size_kb: float
    extraction_method: str


def normalize_text(text: str) -> str:
    """
    Normalize extracted text by removing extra whitespace and newlines.
    Makes the text cleaner for LLM processing and reduces token usage.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Normalized text with consistent spacing
    """
    # Replace multiple spaces/newlines with single space
    normalized = " ".join(text.split())
    return normalized.strip()


def extract_pdf_text(file_content: bytes, file_name: str) -> tuple[str, str]:
    """
    Extract text from PDF file.
    
    Args:
        file_content: Binary content of PDF file
        file_name: Original file name for logging
        
    Returns:
        Tuple of (extracted_text, extraction_method)
        
    Raises:
        HTTPException: If PDF extraction fails
    """
    if PdfReader is None:
        raise HTTPException(
            status_code=500,
            detail="PDF support not installed. Please install PyPDF2."
        )
    
    try:
        # Read PDF from bytes
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the PDF. The file may be scanned/image-based."
            )
        
        logger.info(f"Successfully extracted text from PDF: {file_name} ({len(reader.pages)} pages)")
        return text, "pdf"
        
    except Exception as e:
        logger.error(f"Error extracting PDF {file_name}: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract PDF text: {str(e)}"
        )


def extract_docx_text(file_content: bytes, file_name: str) -> tuple[str, str]:
    """
    Extract text from DOCX file.
    
    Args:
        file_content: Binary content of DOCX file
        file_name: Original file name for logging
        
    Returns:
        Tuple of (extracted_text, extraction_method)
        
    Raises:
        HTTPException: If DOCX extraction fails
    """
    if Document is None:
        raise HTTPException(
            status_code=500,
            detail="DOCX support not installed. Please install python-docx."
        )
    
    try:
        # Read DOCX from bytes
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the DOCX file."
            )
        
        logger.info(f"Successfully extracted text from DOCX: {file_name}")
        return text, "docx"
        
    except Exception as e:
        logger.error(f"Error extracting DOCX {file_name}: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract DOCX text: {str(e)}"
        )


@router.post("/upload", response_model=SyllabusExtractResponse)
async def upload_syllabus(file: UploadFile = File(...)) -> SyllabusExtractResponse:
    """
    Upload and extract text from syllabus document (PDF, DOC, or DOCX).
    
    This endpoint:
    1. Validates file type and size
    2. Extracts raw text from the document
    3. Normalizes the text for LLM processing
    4. Returns the raw text without any interpretation
    
    Accepted formats:
    - PDF (.pdf)
    - DOCX (.docx)
    
    Max file size: 10 MB
    
    Args:
        file: Uploaded document file
        
    Returns:
        SyllabusExtractResponse with extracted raw text
        
    Raises:
        HTTPException: If file validation or extraction fails
    """
    # Validate file extension
    allowed_extensions = {".pdf", ".docx"}
    file_extension = ""
    if file.filename:
        file_extension = "." + file.filename.rsplit(".", 1)[-1].lower()
    
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Accepted formats: {', '.join(allowed_extensions)}"
        )
    
    # Validate file size (10 MB max)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    file_size_bytes = 0
    file_content = b""
    
    try:
        file_content = await file.read()
        file_size_bytes = len(file_content)
        
        if file_size_bytes > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File size exceeds 10 MB limit. Current size: {file_size_bytes / (1024*1024):.2f} MB"
            )
        
        if file_size_bytes == 0:
            raise HTTPException(
                status_code=400,
                detail="File is empty"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reading uploaded file: {str(e)}")
        raise HTTPException(
            status_code=400,
            detail=f"Error reading file: {str(e)}"
        )
    
    # Extract text based on file type
    raw_text = ""
    extraction_method = ""
    
    if file_extension == ".pdf":
        raw_text, extraction_method = extract_pdf_text(file_content, file.filename or "unknown")
    elif file_extension == ".docx":
        raw_text, extraction_method = extract_docx_text(file_content, file.filename or "unknown")
    
    # Normalize the extracted text
    normalized_text = normalize_text(raw_text)
    
    logger.info(
        f"Processed syllabus: {file.filename} | "
        f"Method: {extraction_method} | "
        f"Original chars: {len(raw_text)} | "
        f"Normalized chars: {len(normalized_text)}"
    )
    
    return SyllabusExtractResponse(
        raw_text=normalized_text,
        file_name=file.filename or "unknown",
        file_size_kb=file_size_bytes / 1024,
        extraction_method=extraction_method
    )
